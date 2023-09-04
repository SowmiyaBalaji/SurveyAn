from docx import Document
import collections
import matplotlib.pyplot as plt
from matplotlib.backends.backend_pdf import PdfPages
from docx.shared import Inches
from wordcloud import WordCloud

from pdf2image import convert_from_path
import matplotlib
import pandas as pd
import numpy as np
from itertools import groupby
matplotlib.use('Agg')


def generate_empty_docx(docx_path):
    # Create a new Document object
    doc = Document()

    # Save the document to the specified path
    doc.save(docx_path)


def generate_bar_chart(emotions, frequencies, pdf_path):
    # Create the bar chart using matplotlib
    colors = ["orange", "orange", "orange", 'grey', 'grey', 'grey']
    plt.barh(emotions, frequencies, color=colors)

    # Add labels and title
    for i in range(len(emotions)):
        plt.text(frequencies[i], i, frequencies[i], ha='center')

    plt.xlabel('Emotions')
    plt.ylabel('Frequency')
    plt.title('Emotion Frequency Bar Chart')

    # Save the plot to the PDF file
    with PdfPages(pdf_path) as pdf:
        pdf.savefig()


def generate_pie_chart(emotions, sizes):
    # Create the pie chart
    colors = ['#FFD700', '#0074D9', '#FF4136', '#2ECC40', '#FF851B', '#B10DC9']
    fig, ax = plt.subplots(figsize=(6, 6))
    ax.pie(sizes, labels=emotions, colors=colors,
           autopct='%1.1f%%', startangle=90, shadow=True)

    center_circle = plt.Circle((0, 0), 0.70, fc='white')
    ax.add_patch(center_circle)

    # Equal aspect ratio ensures the pie chart is circular
    ax.axis('equal')

    ax.set_title('Emotion Predictions Distribution')
    ax.legend(emotions, loc='upper right',
              bbox_to_anchor=(1, 0, 0.5, 1), fontsize=12)
    plt.savefig('emotion_frequency_pie.png', dpi=300, bbox_inches='tight')


def generate_wordcloud(y_predrf, df_preproc):

    # Create a DataFrame from the list with column name 'label_num'

    df = pd.DataFrame({'label_num': y_predrf})

    df = pd.concat([df_preproc, df], axis=1)
    print(df.head(10))

    dfjoy = df[df['label_num'] == 0]
    dflove = df[df['label_num'] == 1]
    dfsurprise = df[df['label_num'] == 2]
    dfsadness = df[df['label_num'] == 3]
    dfanger = df[df['label_num'] == 4]
    dffear = df[df['label_num'] == 5]

    text_data = dfjoy['processed_text'].to_string(index=False)

    # Create a WordCloud object and generate the word cloud
    wordcloud = WordCloud(width=800, height=400,
                          background_color='white').generate(text_data)

    # Display the generated word cloud
    plt.figure(figsize=(10, 5))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.title("JOY")
    plt.savefig('wordcloud_joy.png', dpi=300, bbox_inches='tight')

    text_data = dflove['processed_text'].to_string(index=False)

    # Create a WordCloud object and generate the word cloud
    wordcloud = WordCloud(width=800, height=400,
                          background_color='white').generate(text_data)

    # Display the generated word cloud
    plt.figure(figsize=(10, 5))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.title("LOVE")
    plt.savefig('wordcloud_love.png', dpi=300, bbox_inches='tight')

    text_data = dfsurprise['processed_text'].to_string(index=False)

    # Create a WordCloud object and generate the word cloud
    wordcloud = WordCloud(width=800, height=400,
                          background_color='white').generate(text_data)

    # Display the generated word cloud
    plt.figure(figsize=(10, 5))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.title("SURPRISE")
    plt.savefig('wordcloud_surprise.png', dpi=300, bbox_inches='tight')

    text_data = dfsadness['processed_text'].to_string(index=False)

    # Create a WordCloud object and generate the word cloud
    wordcloud = WordCloud(width=800, height=400,
                          background_color='white').generate(text_data)

    # Display the generated word cloud
    plt.figure(figsize=(10, 5))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.title("SAD")
    plt.savefig('wordcloud_sad.png', dpi=300, bbox_inches='tight')

    text_data = dfanger['processed_text'].to_string(index=False)

    # Create a WordCloud object and generate the word cloud
    wordcloud = WordCloud(width=800, height=400,
                          background_color='white').generate(text_data)

    # Display the generated word cloud
    plt.figure(figsize=(10, 5))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.title("ANGER")
    plt.savefig('wordcloud_anger.png', dpi=300, bbox_inches='tight')

    text_data = dffear['processed_text'].to_string(index=False)

    # Create a WordCloud object and generate the word cloud
    wordcloud = WordCloud(width=800, height=400,
                          background_color='white').generate(text_data)

    # Display the generated word cloud
    plt.figure(figsize=(10, 5))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.title("FEAR")
    plt.savefig('wordcloud_fear.png', dpi=300, bbox_inches='tight')


def generate_pdf(y_predrf, df_preproc):

    # Path to the empty Word document file
    empty_docx_path = 'empty_word_document.docx'

    # Generate the empty Word document
    generate_empty_docx(empty_docx_path)

    counter = collections.Counter(y_predrf)
    frequency = [counter[x] for x in sorted(counter.keys())]
    Emotions = ['joy', 'love', 'surprise', 'sadness', 'anger', 'fear']

    # Generate and save the bar chart to a separate PDF file
    generate_bar_chart(Emotions, frequency, 'emotion_frequency_chart.png')
    sizes = [len(list(group)) for key, group in groupby(sorted(y_predrf))]
    generate_pie_chart(Emotions, sizes)

    generate_wordcloud(y_predrf, df_preproc)

    images = convert_from_path('emotion_frequency_chart.pdf')
    png_path = 'emotion_frequency_chart.png'
    images[0].save(png_path)

    # Open the existing Word document and append the bar chart image
    doc = Document(empty_docx_path)
    p = doc.add_paragraph()
    r = p.add_run()
    r.add_text('Emotion Frequency Bar Chart:')
    r.add_picture('emotion_frequency_chart.png', width=Inches(5))
    # Adjust the width as needed
    r.add_text('Emotion Pie Chart:')
    r.add_picture('emotion_frequency_pie.png', width=Inches(5))
    # Adjust the width as needed
    # r.add_text('Happy WordCloud:')
    r.add_picture('wordcloud_joy.png', width=Inches(5))
    # r.add_text('Love WordCloud:')
    r.add_picture('wordcloud_love.png', width=Inches(5))
    # r.add_text('Surprise WordCloud:')
    r.add_picture('wordcloud_surprise.png', width=Inches(5))
    # r.add_text('Sadness WordCloud:')
    r.add_picture('wordcloud_sad.png', width=Inches(5))
    # r.add_text('Anger WordCloud:')
    r.add_picture('wordcloud_anger.png', width=Inches(5))
    # r.add_text('Fear WordCloud:')
    r.add_picture('wordcloud_fear.png', width=Inches(5))
    doc.save('updated_word_file.docx')
